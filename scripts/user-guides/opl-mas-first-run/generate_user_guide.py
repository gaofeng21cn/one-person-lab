from __future__ import annotations

import datetime as dt
import hashlib
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/gaofeng/workspace/one-person-lab")
OUT = ROOT / "dist/user-guides/opl-mas-first-run"
ASSETS = OUT / "assets"
SOURCE = OUT / "source"
VM_SCREENSHOTS = OUT / "vm-screenshots"
VM_ARTIFACTS = OUT / "vm-artifacts"
FINAL_RUN = VM_ARTIFACTS / "final-run"
PDF = OUT / "OPL-MAS-新手首次启动图文教程.pdf"
DOCX = OUT / "OPL-MAS-新手首次启动图文教程.docx"
MD = OUT / "OPL-MAS-新手首次启动图文教程.zh-CN.md"
REFERENCE_DOCX = OUT / "reference-doc.docx"
README = OUT / "README.md"
QA = OUT / "verification-report.json"
RENDER_DIR = ROOT / "tmp/pdfs/opl-mas-first-run/render-pandoc"

RELEASE_TAG = "v26.5.1"
LATEST_RELEASE_URL = "https://github.com/gaofeng21cn/one-person-lab/releases/latest"
LATEST_CHECKED_TAG = "v26.5.2"
LATEST_CHECKED_URL = "https://github.com/gaofeng21cn/one-person-lab/releases/tag/v26.5.2"
RELEASE_URL = "https://github.com/gaofeng21cn/one-person-lab/releases/tag/v26.5.1"
FULL_DMG = "One-Person-Lab-Full-26.5.1-mac-arm64.dmg"
VM_NAME = "opl-userguide-zh-1080p-base"
VM_TYPE = "Tart"
MACOS_GUEST = "26.3"
MACOS_LANGUAGE = "zh-Hans"
MACOS_LOCALE = "zh_CN"
DOCX_FONT = "Microsoft YaHei"
PDF_FONT = "Noto Sans CJK SC"
PDF_SLIDE_SIZE = (2560, 1440)
PDF_SLIDE_DPI = 192.0
PDF_BODY_SIZE = 34
PDF_NOTE_SIZE = 34

FONT_CANDIDATES_REGULAR = [
    "/Users/gaofeng/Library/Fonts/NotoSansCJKsc-Regular.otf",
    "/Users/gaofeng/Library/Fonts/msyh.ttf",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
]
FONT_CANDIDATES_BOLD = [
    "/Users/gaofeng/Library/Fonts/NotoSansCJKsc-Bold.otf",
    "/Users/gaofeng/Library/Fonts/msyhbd.ttf",
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
]

FORBIDDEN_TERMS = [
    "DEMO_KEY",
    "experimental_bearer_token",
    "sk-",
    "OPENAI_API_KEY",
    "CODEX_API_KEY",
    "患者姓名",
    "身份证",
    "独立安装包",
    "独立 Release",
    "单独安装 MAS",
]

DOCX_IMAGE_HEIGHT = "2.75in"


@dataclass(frozen=True)
class AssetSpec:
    source: str
    output: str
    alt: str
    crop: tuple[int, int, int, int] | None = None
    redactions: tuple[tuple[int, int, int, int], ...] = ()


@dataclass(frozen=True)
class Step:
    title: str
    body: str
    asset_key: str
    notes: tuple[str, ...]


ASSET_SPECS = [
    AssetSpec("01-download-release.png", "01-download-release.png", "Release 下载页面"),
    AssetSpec("02-install-dmg.png", "02-install-dmg.png", "DMG 安装窗口", (720, 320, 1800, 1148)),
    AssetSpec(
        "03-first-launch.png",
        "03-codex-config-needed.png",
        "首次启动需要配置 Codex",
        (240, 160, 3600, 1960),
        ((1060, 1320, 2600, 1545),),
    ),
    AssetSpec("04-codex-configured.png", "04-first-run-checking.png", "首次启动环境检查", (240, 160, 3600, 1960)),
    AssetSpec("05-after-prepare.png", "05-opl-ready-research-entry.png", "OPL 主界面和科研入口", (240, 160, 3600, 1960)),
    AssetSpec("06-data-folder-raw-data.png", "06-research-data-folder.png", "示例专病 workspace 和 raw_data 文件夹", (240, 180, 2200, 1680)),
    AssetSpec("07-first-task-prompt-2.png", "07-first-research-entry.png", "首次科研任务输入区", (240, 160, 3600, 1960)),
    AssetSpec("08-opl-status.png", "08-opl-runtime-status.png", "OPL 运行状态页面", (240, 160, 3600, 1960)),
]

STEPS = [
    Step(
        "1. 下载 One Person Lab",
        "访问 One Person Lab 最新 Release 页面，下载最新 macOS Apple Silicon DMG。首次安装或干净机器建议下载 Full 版本。",
        "01-download-release.png",
        (
            f"最新版本页面：[{LATEST_RELEASE_URL}]({LATEST_RELEASE_URL})。",
            "Full 首次安装包包含 MAS/Hermes/MDS 等运行 payload，适合首次安装、干净机器或需要完整离线首启资产的用户。",
            "标准 mac-arm64 DMG 体积更小，适合已经安装 OPL、后续通过 App 内更新或标准包更新的用户。",
            "首次下载建议选择 Full / mac-arm64 DMG；下载前确认来源页面并核对 Release 页面校验信息。",
        ),
    ),
    Step(
        "2. 安装 App",
        "打开 DMG，将 One Person Lab 拖入 Applications。首次打开如出现 macOS 安全提示，按系统提示确认。",
        "02-install-dmg.png",
        (
            "安装完成后从 Applications 启动 App。",
            "不要长期在 DMG 挂载窗口内运行 App。",
        ),
    ),
    Step(
        "3. 配置 Codex 权限",
        "首次启动如果要求 API Key 或 Codex 权限，统一联系 gflabtoken 管理员开通。",
        "03-codex-config-needed.png",
        (
            "管理员开通后，按管理员给出的方式完成配置。",
            "不要把密钥截图、转发或写入研究数据目录。",
        ),
    ),
    Step(
        "4. 等待首次环境检查",
        "OPL 会检查 Codex、模块、skills 和本机运行环境。等待状态进入可继续阶段。",
        "04-first-run-checking.png",
        (
            "首启准备可能需要几分钟。",
            "遇到阻塞时先阅读界面提示，再联系技术支持处理。",
        ),
    ),
    Step(
        "5. 进入科研入口",
        "准备完成后，在主界面选择“科研”，进入 Research Foundry / Med Auto Science 工作流。",
        "05-opl-ready-research-entry.png",
        (
            "MAS 通过 OPL 内的 Research Foundry / Med Auto Science 入口使用。",
            "用户不需要另行获取 MAS 分发资产。",
        ),
    ),
    Step(
        "6. 准备研究数据目录",
        "建议按一个病种或稳定研究主题新建本地 workspace，把原始或脱敏材料集中放入 `raw_data/`。",
        "06-research-data-folder.png",
        (
            "新手首启阶段不需要手工建立 MAS 内部目录结构。",
            "MAS 会按内置规则初始化 workspace，并吸收整理 `raw_data/` 中的材料到可用形式。",
            "患者数据需先脱敏，并遵守本机构数据管理要求。",
        ),
    ),
    Step(
        "7. 发起首次科研任务",
        "第一条任务可直接用自然语言描述，让 MAS 先判断研究方向、证据缺口和下一步。",
        "07-first-research-entry.png",
        (
            "示例提示词：我有一批肺结节随访数据，专病 workspace 在“肺结节真实世界研究”，原始材料在 `raw_data/`，请先判断最值得推进的研究问题，并说明还缺哪些证据，目标是形成一篇可投稿论文。",
        ),
    ),
    Step(
        "8. 查看进度与结果",
        "任务启动后，重点查看当前阶段、阻塞项、下一步和产物位置。",
        "08-opl-runtime-status.png",
        (
            "看到需要人工确认的项目时，由研究者或 PI 判断是否继续。",
            "投稿前最终科学判断、伦理合规和署名安排仍由研究团队负责。",
        ),
    ),
]


def _font_path(candidates: list[str]) -> str:
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return candidates[0]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(_font_path(FONT_CANDIDATES_BOLD if bold else FONT_CANDIDATES_REGULAR), size)


def wrap(text: str, size: int, width: int, bold: bool = False) -> list[str]:
    f = font(size, bold=bold)
    lines: list[str] = []
    current = ""
    for ch in text:
        if ch == "\n":
            lines.append(current)
            current = ""
            continue
        candidate = current + ch
        if f.getlength(candidate) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    content: str,
    size: int,
    fill: str = "#172033",
    bold: bool = False,
) -> None:
    draw.text(xy, content, font=font(size, bold=bold), fill=fill)


def multiline(
    draw: ImageDraw.ImageDraw,
    x: int,
    y: int,
    content: str,
    size: int,
    width: int,
    fill: str = "#314057",
    line_gap: int = 10,
    bold: bool = False,
) -> int:
    for line in wrap(content, size, width, bold=bold):
        draw.text((x, y), line, font=font(size, bold=bold), fill=fill)
        y += size + line_gap
    return y


def rounded(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    radius: int,
    fill: str,
    outline: str | None = None,
    width: int = 1,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def rounded_outline(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    radius: int,
    outline: str,
    width: int = 1,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=None, outline=outline, width=width)


def release_sha256() -> str:
    dmg = SOURCE / FULL_DMG
    return hashlib.sha256(dmg.read_bytes()).hexdigest() if dmg.exists() else ""


def bullet_text(note: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", note).replace("`", "")


def make_assets() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    assets: dict[str, Path] = {}
    for old_asset in ASSETS.glob("*.png"):
        old_asset.unlink()

    for spec in ASSET_SPECS:
        src = VM_SCREENSHOTS / spec.source
        dst = ASSETS / spec.output
        if not src.exists():
            raise FileNotFoundError(src)
        if spec.crop:
            with Image.open(src) as image:
                output = image.crop(spec.crop).convert("RGBA")
        else:
            with Image.open(src) as image:
                output = image.convert("RGBA")
        if spec.redactions:
            draw = ImageDraw.Draw(output)
            for box in spec.redactions:
                draw.rectangle(box, fill="#fbfbfb")
        output.save(dst)
        assets[spec.output] = dst
    return assets


def make_markdown(assets: dict[str, Path]) -> None:
    lines = [
        "---",
        "title-meta: OPL + MAS 新手首次启动图文教程",
        "lang: zh-Hans",
        "---",
        "",
        "# OPL + MAS 新手首次启动图文教程",
        "",
        "适用对象：医生、PI、课题负责人；不要求计算机基础。本文以 macOS App 首次启动为主线，说明如何下载、安装、配置 One Person Lab，并通过 Research Foundry / Med Auto Science 发起首次科研任务。",
        "",
        "> 涉及 Codex API Key 或 Codex 权限配置时，请联系 gflabtoken 管理员开通。不要自行购买、复制来源不明的密钥，或把密钥写入研究数据目录。",
        "",
        "## 准备清单",
        "",
        "- 一台 Apple Silicon Mac 或可运行 macOS App 的 Mac。",
        "- 稳定网络，用于下载 One Person Lab 和完成首次环境检查。",
        "- gflabtoken 开通状态；涉及 Codex 权限时请联系 gflabtoken 管理员。",
        "- 本地研究数据文件夹，数据需完成脱敏并符合本机构数据管理要求。",
        "- 变量说明、纳排标准、终点定义、统计计划、参考文献或已有草稿；可以先放入专病 workspace 的 `raw_data/`。",
        "",
    ]

    for step in STEPS:
        image = assets[step.asset_key].relative_to(OUT)
        lines.extend(
            [
                f"## {step.title}",
                "",
                step.body,
                "",
                f"![]({image}){{height={DOCX_IMAGE_HEIGHT}}}",
                "",
            ]
        )
        for note in step.notes:
            lines.append(f"- {note}")
        lines.append("")

    lines.extend(
        [
            "## 常见问题",
            "",
            "- 下载失败：换网络后重试，或请技术支持人员确认 GitHub Release 是否可访问。",
            "- 打不开 App：确认已拖入 Applications，并按 macOS 安全提示允许打开。",
            "- Codex 未配置：联系 gflabtoken 管理员开通。",
            "- 模块未就绪：在 App 的环境管理中重新检查，确认 OPL 完整安装资产与本机网络状态。",
            "- 数据路径看不到：确认选择的是本机可访问的专病 workspace，或能看到其中的 `raw_data/`。",
            "- 任务启动后不知道看哪里：查看运行状态页的当前阶段、下一步和需要人工确认的项目。",
            "",
        ]
    )
    MD.write_text("\n".join(lines), encoding="utf-8")



def make_reference_docx() -> None:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.62)
        section.right_margin = Inches(0.62)

    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = DOCX_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), DOCX_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOCX_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), DOCX_FONT)
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)

    doc.add_heading("OPL + MAS 新手首次启动图文教程", 0)
    doc.add_paragraph("本文件仅作为 Pandoc reference-doc，用于控制页边距和基础中文字体。")
    doc.save(REFERENCE_DOCX)


def run_command(command: list[str]) -> None:
    subprocess.run(command, cwd=OUT, check=True)


def pandoc_version() -> str:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return ""
    result = subprocess.run([pandoc, "--version"], check=True, text=True, capture_output=True)
    return result.stdout.splitlines()[0]


def xelatex_version() -> str:
    xelatex = shutil.which("xelatex")
    if not xelatex:
        return ""
    result = subprocess.run([xelatex, "--version"], check=True, text=True, capture_output=True)
    return result.stdout.splitlines()[0]


def paste_contained(canvas: Image.Image, image: Image.Image, box: tuple[int, int, int, int]) -> tuple[int, int, int, int]:
    left, top, right, bottom = box
    width = right - left
    height = bottom - top
    scale = min(width / image.width, height / image.height)
    new_size = (int(image.width * scale), int(image.height * scale))
    resized = image.convert("RGB").resize(new_size, Image.Resampling.LANCZOS)
    x = left + (width - new_size[0]) // 2
    y = top + (height - new_size[1]) // 2
    canvas.paste(resized, (x, y))
    return (x, y, x + new_size[0], y + new_size[1])


def draw_bullets(
    draw: ImageDraw.ImageDraw,
    notes: tuple[str, ...],
    x: int,
    y: int,
    width: int,
    size: int,
    line_gap: int = 10,
    item_gap: int = 10,
) -> int:
    for note in notes:
        clean = bullet_text(note)
        text(draw, (x, y), "•", size, fill="#2f5f9d", bold=True)
        y = multiline(draw, x + 42, y, clean, size, width - 42, fill="#25324a", line_gap=line_gap)
        y += item_gap
    return y


def slide_canvas(title_text: str, subtitle: str | None = None) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    canvas = Image.new("RGB", PDF_SLIDE_SIZE, "#f6f8fb")
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, PDF_SLIDE_SIZE[0], 104), fill="#ffffff")
    text(draw, (96, 35), title_text, 42, fill="#244f85", bold=True)
    if subtitle:
        multiline(draw, 1420, 34, subtitle, 22, 990, fill="#53657d", line_gap=9)
    draw.line((96, 104, 2464, 104), fill="#d8e2f0", width=3)
    return canvas, draw


def make_cover_slide() -> Image.Image:
    canvas, draw = slide_canvas("OPL + MAS 新手首次启动图文教程", "面向医生、PI 和课题负责人；以 macOS App 首次启动路径为主线。")
    text(draw, (130, 165), "准备清单", 38, fill="#244f85", bold=True)
    checklist = (
        "一台 Apple Silicon Mac 或可运行 macOS App 的 Mac。",
        "稳定网络，用于下载 One Person Lab 和完成首次环境检查。",
        "gflabtoken 开通状态；涉及 Codex 权限时请联系 gflabtoken 管理员。",
        "本地专病 workspace；患者数据需完成脱敏并符合本机构数据管理要求。",
        "变量说明、纳排标准、终点定义、统计计划、参考文献或已有草稿，可先放入 raw_data/。",
    )
    draw_bullets(draw, checklist, 150, 245, 2180, 31)
    rounded(draw, (130, 760, 2430, 1028), 22, "#ffffff", outline="#d7e2ef", width=2)
    multiline(
        draw,
        175,
        805,
        "涉及 Codex API Key 或 Codex 权限配置时，请联系 gflabtoken 管理员开通。不要自行购买、复制来源不明的密钥，或把密钥写入研究数据目录。",
        34,
        2130,
        fill="#223149",
        line_gap=16,
    )
    text(draw, (130, 1262), "截图环境：中文 macOS，Retina/HiDPI 开启，等效 1920x1080，原始截图 3840x2160。", 24, fill="#5c6c80")
    return canvas


def make_step_slide(step: Step, asset: Path, page_no: int) -> Image.Image:
    canvas, draw = slide_canvas(step.title)
    multiline(draw, 112, 132, step.body, PDF_BODY_SIZE, 2250, fill="#162033", line_gap=16)
    with Image.open(asset) as image:
        pasted_box = paste_contained(canvas, image, (112, 250, 2448, 1120))
    draw.rounded_rectangle(
        (pasted_box[0] - 8, pasted_box[1] - 8, pasted_box[2] + 8, pasted_box[3] + 8),
        radius=18,
        outline="#d3deea",
        width=2,
    )
    draw_bullets(draw, step.notes, 130, 1165, 2220, PDF_NOTE_SIZE, line_gap=12, item_gap=10)
    text(draw, (2380, 1338), str(page_no), 22, fill="#8b99aa")
    return canvas


def make_faq_slide(page_no: int) -> Image.Image:
    canvas, draw = slide_canvas("常见问题")
    faqs = (
        "下载失败：换网络后重试，或请技术支持人员确认 GitHub Release 是否可访问。",
        "打不开 App：确认已拖入 Applications，并按 macOS 安全提示允许打开。",
        "Codex 未配置：联系 gflabtoken 管理员开通。",
        "模块未就绪：在 App 的环境管理中重新检查，确认 OPL 完整安装资产与本机网络状态。",
        "数据路径看不到：确认选择的是本机可访问的专病 workspace，或能看到其中的 raw_data/。",
        "任务启动后不知道看哪里：查看运行状态页的当前阶段、下一步和需要人工确认的项目。",
    )
    draw_bullets(draw, faqs, 150, 175, 2180, PDF_NOTE_SIZE, line_gap=16, item_gap=24)
    text(draw, (2380, 1338), str(page_no), 22, fill="#8b99aa")
    return canvas


def make_pdf(assets: dict[str, Path]) -> None:
    slides = [make_cover_slide()]
    for index, step in enumerate(STEPS, start=1):
        slides.append(make_step_slide(step, assets[step.asset_key], index + 1))
    slides.append(make_faq_slide(len(slides) + 1))
    slides[0].save(
        PDF,
        "PDF",
        resolution=PDF_SLIDE_DPI,
        save_all=True,
        append_images=slides[1:],
    )


def normalize_docx_fonts(path: Path) -> None:
    replacements = {
        "ＭＳ 明朝": DOCX_FONT,
        "ＭＳ ゴシック": DOCX_FONT,
        "宋体": DOCX_FONT,
        "SimSun": DOCX_FONT,
        "MS Mincho": DOCX_FONT,
        "MS Gothic": DOCX_FONT,
    }
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_path = Path(tmp_file.name)
    try:
        with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in {"word/styles.xml", "word/theme/theme1.xml", "word/fontTable.xml", "word/settings.xml"}:
                    text_data = data.decode("utf-8")
                    for old, new in replacements.items():
                        text_data = text_data.replace(old, new)
                    data = text_data.encode("utf-8")
                zout.writestr(item, data)
        tmp_path.replace(path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def make_documents() -> None:
    make_reference_docx()
    common = [
        "--from", "markdown+link_attributes",
        "--resource-path", str(OUT),
        "--metadata", "lang=zh-Hans",
        "--metadata", "title-meta=OPL + MAS 新手首次启动图文教程",
    ]
    run_command([
        "pandoc",
        str(MD),
        "-o",
        str(DOCX),
        "--reference-doc",
        str(REFERENCE_DOCX),
        *common,
    ])
    normalize_docx_fonts(DOCX)

def read_display_proof() -> str:
    proof = FINAL_RUN / "display-retina-effective-1080p.txt"
    return proof.read_text(encoding="utf-8").strip() if proof.exists() else ""


def make_readme(assets: dict[str, Path]) -> None:
    sha = release_sha256()
    display_proof = read_display_proof()
    asset_lines = "\n".join(f"- `assets/{path.name}`" for path in assets.values())
    content = f"""# OPL + MAS 新手首次启动图文教程

## 交付物

- `OPL-MAS-新手首次启动图文教程.zh-CN.md`
- `OPL-MAS-新手首次启动图文教程.docx`
- `OPL-MAS-新手首次启动图文教程.pdf`
- `assets/*.png`
- `verification-report.json`

## Release 信息

- 最新 Release 通用页面: {LATEST_RELEASE_URL}
- 最后检查的 latest Release: `{LATEST_CHECKED_TAG}`
- 最后检查的 latest Release 页面: {LATEST_CHECKED_URL}
- 本轮截图与首启验证使用的 Release: `{RELEASE_TAG}`
- 本轮截图与首启验证页面: {RELEASE_URL}
- 本轮截图与首启验证使用的首次安装资产: `{FULL_DMG}`
- SHA256: `{sha}`

## VM 与截图环境

- VM 类型: `{VM_TYPE}`
- VM 名称: `{VM_NAME}`
- macOS guest: `{MACOS_GUEST}`
- 系统语言: `{MACOS_LANGUAGE}`
- Locale: `{MACOS_LOCALE}`
- 显示模式: Retina/HiDPI 开启后的等效 `1920x1080`
- 截图像素: `3840x2160`

显示验证记录:

```text
{display_proof}
```

截图采集方式:

1. 在 Tart guest 内通过 `CGDisplayCopyAllDisplayModes(..., kCGDisplayShowDuplicateLowResolutionModes=true)` 找到 `1920x1080` logical、`3840x2160` pixel 的 HiDPI mode。
2. 使用 `CGConfigureDisplayWithDisplayMode` 切换到该模式。
3. 通过 guest 内 `screencapture -x` 采集 PNG。
4. 用 `sips` / PIL 验证截图像素为 `3840x2160`。

## 截图资源

{asset_lines}

## 首启验证工件

- `vm-artifacts/final-run/first-run.jsonl`
- `vm-artifacts/final-run/opl-system-initialize.json`
- `vm-artifacts/final-run/opl-modules.json`
- `vm-artifacts/final-run/display-retina-effective-1080p.txt`

首启结果摘要:

- `first-run.jsonl` 记录了从 `codex-config-needed` 到 `prepared` 的 GUI 首启状态。
- `opl system initialize --json` 返回 `setup_phase=review`，`blocking_items=[]`。
- `opl modules --json` 用于确认模块枚举与状态面可读取。

## 文档生成

本目录的 Word 由 Markdown 源稿通过 Pandoc 生成；PDF 使用同一份步骤内容与截图资源生成横版图文页。

```bash
python3 scripts/user-guides/opl-mas-first-run/generate_user_guide.py
```

Word 核心命令:

```bash
pandoc dist/user-guides/opl-mas-first-run/OPL-MAS-新手首次启动图文教程.zh-CN.md \
  -o dist/user-guides/opl-mas-first-run/OPL-MAS-新手首次启动图文教程.docx \
  --reference-doc dist/user-guides/opl-mas-first-run/reference-doc.docx \
  --resource-path dist/user-guides/opl-mas-first-run
```

生成工具:

- Pandoc: `{pandoc_version()}`
- DOCX 中文字体: `{DOCX_FONT}`
- PDF 版式: 横版图文页，一步骤一页

## QA

```bash
mkdir -p tmp/pdfs/opl-mas-first-run/render-pandoc
pdftoppm -png -r 150 dist/user-guides/opl-mas-first-run/OPL-MAS-新手首次启动图文教程.pdf tmp/pdfs/opl-mas-first-run/render-pandoc/page
officecli validate dist/user-guides/opl-mas-first-run/OPL-MAS-新手首次启动图文教程.docx
officecli view dist/user-guides/opl-mas-first-run/OPL-MAS-新手首次启动图文教程.docx outline
```

维护入口:

- `scripts/user-guides/opl-mas-first-run/README.md`
- `scripts/user-guides/opl-mas-first-run/generate_user_guide.py`

安全处理:

- 未使用真实 API key、token、账号、患者身份信息或未脱敏机构数据。
- 教程主图从原始 Retina 截图裁切；涉及本地日志路径的区域已遮挡。
- 安装 App 截图保持 DMG 默认窗口尺寸，不通过拉大窗口制造居中假象。
- Codex 配置统一表述为“联系 gflabtoken 管理员开通”。
- MAS 表述为 OPL 内的 Research Foundry / Med Auto Science 入口。
"""
    README.write_text(content, encoding="utf-8")


def image_dimensions(path: Path) -> tuple[int, int]:
    with Image.open(path) as image:
        return image.size


def display_checks(display_proof: str) -> dict[str, bool]:
    return {
        "logical_1920x1080": "current_mode=1920x1080" in display_proof and "frame_points=1920x1080" in display_proof,
        "pixel_3840x2160": "pixel=3840x2160" in display_proof,
        "backing_scale_2": "backing_scale=2.0" in display_proof,
    }


def capture_command(command: list[str]) -> dict[str, object]:
    executable = shutil.which(command[0])
    if not executable:
        return {
            "command": command,
            "available": False,
            "returncode": None,
            "stdout": "",
            "stderr": f"{command[0]} not found",
        }
    completed = subprocess.run(
        [executable, *command[1:]],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    return {
        "command": command,
        "available": True,
        "returncode": completed.returncode,
        "stdout": completed.stdout[-6000:],
        "stderr": completed.stderr[-6000:],
    }


def pdfinfo_check() -> dict[str, object]:
    result = capture_command(["pdfinfo", str(PDF)])
    info: dict[str, str] = {}
    if result["stdout"]:
        for line in str(result["stdout"]).splitlines():
            if ":" not in line:
                continue
            key, value = line.split(":", 1)
            info[key.strip()] = value.strip()
    result["parsed"] = info
    result["pages"] = int(info["Pages"]) if info.get("Pages", "").isdigit() else None
    return result


def make_contact_sheet(pages: list[Path]) -> None:
    if not pages:
        return
    thumb_w = 360
    thumbs: list[Image.Image] = []
    for page in pages:
        with Image.open(page) as image:
            rgb = image.convert("RGB")
            scale = thumb_w / rgb.width
            thumb = rgb.resize((thumb_w, int(rgb.height * scale)))
        canvas = Image.new("RGB", (thumb_w, thumb.height + 34), "white")
        canvas.paste(thumb, (0, 0))
        drawer = ImageDraw.Draw(canvas)
        drawer.text((4, thumb.height + 8), page.name, fill="black")
        thumbs.append(canvas)
    cols = 4
    rows = (len(thumbs) + cols - 1) // cols
    cell_h = max(thumb.height for thumb in thumbs)
    sheet = Image.new("RGB", (cols * thumb_w, rows * cell_h), "#f3f4f6")
    for index, thumb in enumerate(thumbs):
        sheet.paste(thumb, ((index % cols) * thumb_w, (index // cols) * cell_h))
    sheet.save(RENDER_DIR / "contact-sheet.png")


def render_pdf_check() -> dict[str, object]:
    pdftoppm = shutil.which("pdftoppm")
    if not pdftoppm:
        return {
            "command": ["pdftoppm", "-png", "-r", "150", str(PDF), str(RENDER_DIR / "page")],
            "available": False,
            "returncode": None,
            "page_count": 0,
            "page_dimensions": {},
            "contact_sheet": "",
        }
    if RENDER_DIR.exists():
        shutil.rmtree(RENDER_DIR)
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    completed = subprocess.run(
        [pdftoppm, "-png", "-r", "150", str(PDF), str(RENDER_DIR / "page")],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    pages = sorted(RENDER_DIR.glob("page-*.png"))
    page_dimensions = {page.name: image_dimensions(page) for page in pages}
    make_contact_sheet(pages)
    return {
        "command": ["pdftoppm", "-png", "-r", "150", str(PDF), str(RENDER_DIR / "page")],
        "available": True,
        "returncode": completed.returncode,
        "stdout": completed.stdout[-6000:],
        "stderr": completed.stderr[-6000:],
        "page_count": len(pages),
        "page_dimensions": page_dimensions,
        "contact_sheet": str(RENDER_DIR / "contact-sheet.png"),
    }


def officecli_docx_checks() -> dict[str, object]:
    return {
        "validate": capture_command(["officecli", "validate", str(DOCX)]),
        "outline": capture_command(["officecli", "view", str(DOCX), "outline"]),
    }


def scan_forbidden(paths: list[Path]) -> list[dict[str, str]]:
    hits: list[dict[str, str]] = []
    pattern = re.compile("|".join(re.escape(term) for term in FORBIDDEN_TERMS))
    for path in paths:
        if not path.exists() or path.is_dir():
            continue
        if path.suffix.lower() in {".png", ".pdf", ".dmg", ".zip"}:
            continue
        for line_no, line in enumerate(path.read_text(encoding="utf-8", errors="ignore").splitlines(), start=1):
            for match in pattern.finditer(line):
                hits.append({"path": str(path.relative_to(OUT)), "line": str(line_no), "term": match.group(0)})
    return hits


def make_qa(assets: dict[str, Path]) -> None:
    display_proof = read_display_proof()
    source_dimensions = {
        spec.output: image_dimensions(VM_SCREENSHOTS / spec.source)
        for spec in ASSET_SPECS
    }
    asset_dimensions = {
        path.name: image_dimensions(path)
        for path in assets.values()
    }
    source_retina_ok = all(dim == (3840, 2160) for dim in source_dimensions.values())
    small_default_window_assets = {
        name: dim
        for name, dim in asset_dimensions.items()
        if name in {"02-install-dmg.png"}
    }
    derived_assets_ok = all(
        width >= 1920 and height >= 1080
        for name, (width, height) in asset_dimensions.items()
        if name not in small_default_window_assets
    )
    forbidden_hits = scan_forbidden(
        [
            MD,
            README,
            FINAL_RUN / "first-run.jsonl",
            FINAL_RUN / "opl-system-initialize.json",
            FINAL_RUN / "opl-modules.json",
            FINAL_RUN / "display-retina-effective-1080p.txt",
        ]
    )
    checks = {
        "generated_at": dt.datetime.now().isoformat(),
        "release": RELEASE_TAG,
        "latest_release_checked_tag": LATEST_CHECKED_TAG,
        "latest_release_checked_url": LATEST_CHECKED_URL,
        "full_dmg": FULL_DMG,
        "full_dmg_sha256": release_sha256(),
        "pdf": str(PDF),
        "docx": str(DOCX),
        "markdown": str(MD),
        "readme": str(README),
        "pandoc_version": pandoc_version(),
        "xelatex_version": xelatex_version(),
        "asset_count": len(assets),
        "pdf_exists": PDF.exists(),
        "docx_exists": DOCX.exists(),
        "markdown_exists": MD.exists(),
        "readme_exists": README.exists(),
        "display_proof": display_proof,
        "display_checks": display_checks(display_proof),
        "source_screenshot_dimensions": source_dimensions,
        "asset_dimensions": asset_dimensions,
        "source_screenshots_all_3840x2160": source_retina_ok,
        "derived_assets_min_1920x1080_except_default_window_crops": derived_assets_ok,
        "default_window_crops": small_default_window_assets,
        "pdfinfo": pdfinfo_check(),
        "pdf_render_150dpi": render_pdf_check(),
        "docx_officecli": officecli_docx_checks(),
        "forbidden_hits": forbidden_hits,
        "forbidden_terms_absent": not forbidden_hits,
        "notes": [
            "Tutorial assets are copied or cropped from real Chinese macOS VM screenshots.",
            "The VM display was set to Retina/HiDPI effective 1920x1080 with backing_scale=2.0.",
            "Most cropped derived assets keep at least 1920x1080 pixels for PDF clarity.",
            "The DMG install screenshot is intentionally cropped to the natural default window rather than enlarged.",
        ],
    }
    QA.write_text(json.dumps(checks, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    SOURCE.mkdir(parents=True, exist_ok=True)
    assets = make_assets()
    make_markdown(assets)
    make_documents()
    make_pdf(assets)
    make_readme(assets)
    make_qa(assets)
    print(PDF)
    print(DOCX)
    print(MD)
    print(README)
    print(QA)


if __name__ == "__main__":
    main()
